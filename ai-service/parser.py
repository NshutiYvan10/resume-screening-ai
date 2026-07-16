"""
Resume text extraction module
Supports PDF, DOCX, DOC, and TXT formats
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """
    Extract raw text from resume files
    
    Supported formats:
    - PDF: uses pdfminer.six
    - DOCX: uses python-docx
    - DOC: fallback to python-docx or textract
    - TXT: direct read
    
    Returns cleaned string with normalized whitespace
    """
    if not os.path.exists(file_path):
        logger.warning(f"File not found: {file_path}")
        return ""
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            return _extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return _extract_docx(file_path)
        elif file_ext == '.txt':
            return _extract_txt(file_path)
        else:
            logger.warning(f"Unsupported file format: {file_ext}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfminer.six"""
    try:
        from pdfminer.high_level import extract_text as pdf_extract
        
        text = pdf_extract(file_path)
        return _clean_text(text)
    except ImportError:
        logger.error("pdfminer.six not installed. Install with: pip install pdfminer.six")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return ""


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX/DOC using python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = "\n".join(text_parts)
        return _clean_text(text)
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return ""


def _extract_txt(file_path: str) -> str:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return _clean_text(text)
    except Exception as e:
        logger.error(f"TXT extraction error: {str(e)}")
        return ""


def _clean_text(text: str) -> str:
    """
    Clean extracted text while PRESERVING line structure.

    Line structure is load-bearing: section detection (education vs experience),
    name extraction and date-range attribution all depend on it. Only intra-line
    whitespace is collapsed; empty lines are dropped.
    """
    if not text:
        return ""

    lines = []
    for line in text.split("\n"):
        # collapse tabs/multiple spaces within the line
        line = ' '.join(line.split())
        # strip non-printable characters
        line = "".join(ch for ch in line if ch.isprintable())
        if line:
            lines.append(line)

    return "\n".join(lines).strip()